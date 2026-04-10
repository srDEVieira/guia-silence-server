from __future__ import annotations

from pathlib import Path
import tempfile
import threading
from typing import Iterable
import zipfile
import atexit

from docx import Document
from lxml import etree
import time

try:
    import xlrd
except ModuleNotFoundError:
    xlrd = None

try:
    import win32com.client as win32
except ModuleNotFoundError:
    win32 = None

try:
    import win32print
except ModuleNotFoundError:
    win32print = None


ItemRow = tuple[str, str]
_WORD_WARMED = False
_WORD_APP = None
_WORD_APP_LOCK = threading.Lock()


def _ensure_word_automation() -> None:
    if win32 is None:
        raise RuntimeError(
            "A impressao pelo Word exige a biblioteca pywin32 instalada na .venv. "
            "Rode: .\\.venv\\Scripts\\python.exe -m pip install -r requirements.txt"
        )


def warmup_word_automation() -> None:
    global _WORD_WARMED

    if _WORD_WARMED or win32 is None:
        return

    try:
        app = _get_word_application()
        _ = app.Version
        _WORD_WARMED = True
    except Exception:
        return


def _get_word_application():
    global _WORD_APP

    _ensure_word_automation()

    with _WORD_APP_LOCK:
        if _WORD_APP is not None:
            try:
                _ = _WORD_APP.Version
                return _WORD_APP
            except Exception:
                _WORD_APP = None

        app = win32.DispatchEx("Word.Application")
        app.Visible = False
        app.DisplayAlerts = 0
        _WORD_APP = app
        return app


def shutdown_word_automation() -> None:
    global _WORD_APP

    with _WORD_APP_LOCK:
        if _WORD_APP is None:
            return
        try:
            try:
                _WORD_APP.NormalTemplate.Saved = True
            except Exception:
                pass
            _WORD_APP.Quit()
        except Exception:
            pass
        finally:
            _WORD_APP = None


atexit.register(shutdown_word_automation)


def load_inventory_lookup(spreadsheet_path: Path) -> dict[str, str]:
    if not spreadsheet_path.is_file():
        return {}
    lookup: dict[str, str] = {}

    if xlrd is not None:
        workbook = xlrd.open_workbook(spreadsheet_path)
        sheet = workbook.sheet_by_index(0)

        for row_index in range(1, sheet.nrows):
            chapa_key = normalize_patrimony(sheet.cell_value(row_index, 0))
            description_text = str(sheet.cell_value(row_index, 1)).strip()

            if chapa_key and description_text and chapa_key not in lookup:
                lookup[chapa_key] = description_text

        return lookup

    _ensure_word_automation()

    excel = win32.DispatchEx("Excel.Application")
    excel.Visible = False
    excel.DisplayAlerts = False
    workbook = None

    try:
        workbook = excel.Workbooks.Open(str(spreadsheet_path.resolve()))
        worksheet = workbook.Worksheets(1)
        used_range = worksheet.UsedRange
        values = used_range.Value

        if not values:
            return {}

        rows = values if isinstance(values, tuple) else (values,)
        for row in rows[1:]:
            if not isinstance(row, tuple) or len(row) < 2:
                continue

            chapa_key = normalize_patrimony(row[0])
            description_text = "" if row[1] is None else str(row[1]).strip()

            if chapa_key and description_text and chapa_key not in lookup:
                lookup[chapa_key] = description_text
    finally:
        if workbook is not None:
            try:
                workbook.Close(False)
            except Exception:
                pass
        excel.Quit()

    return lookup


def normalize_patrimony(value: object) -> str:
    if value is None:
        return ""

    text = str(value).strip()
    if not text:
        return ""

    if text.endswith(".0"):
        text = text[:-2]

    return "".join(character for character in text if character.isalnum()).casefold()


def list_printers() -> list[str]:
    if win32print is None:
        return []

    printers: list[str] = []
    flag_sets = [
        win32print.PRINTER_ENUM_LOCAL | win32print.PRINTER_ENUM_CONNECTIONS,
        win32print.PRINTER_ENUM_LOCAL,
        win32print.PRINTER_ENUM_CONNECTIONS,
    ]

    for flags in flag_sets:
        try:
            for item in win32print.EnumPrinters(flags):
                name = str(item[2]).strip()
                if name and name not in printers:
                    printers.append(name)
        except Exception:
            continue

    return sorted(printers, key=str.casefold)


def get_default_printer_name() -> str:
    if win32print is None:
        return ""
    try:
        return win32print.GetDefaultPrinter()
    except Exception:
        return ""


def print_docx(document_path: Path, printer_name: str, copies: int = 1) -> None:
    print_docx_batch([document_path], printer_name, copies=copies)


def print_docx_batch(document_paths: Iterable[Path], printer_name: str, copies: int = 1) -> None:
    paths = [Path(path).resolve() for path in document_paths]
    if not paths:
        return

    app = _get_word_application()
    previous_printer = None

    try:
        previous_printer = app.ActivePrinter
    except Exception:
        previous_printer = None

    try:
        if printer_name:
            app.ActivePrinter = printer_name

        for document_path in paths:
            document = None
            try:
                document = app.Documents.Open(str(document_path), ReadOnly=True, AddToRecentFiles=False)
                document.PrintOut(Background=True, Copies=max(1, int(copies)))
            finally:
                if document is not None:
                    document.Close(SaveChanges=False)

        time.sleep(0.01)
    finally:
        try:
            app.NormalTemplate.Saved = True
        except Exception:
            pass
        if previous_printer:
            try:
                app.ActivePrinter = previous_printer
            except Exception:
                pass


def print_docx_with_dialog(document_path: Path) -> None:
    _ensure_word_automation()

    app = win32.DispatchEx("Word.Application")
    app.Visible = True
    app.DisplayAlerts = 0
    document = None

    try:
        document = app.Documents.Open(str(document_path.resolve()), ReadOnly=True, AddToRecentFiles=False)
        document.Activate()
        app.Activate()
        app.Dialogs(88).Show()
        time.sleep(0.5)
    finally:
        if document is not None:
            try:
                document.Close(SaveChanges=False)
            except Exception:
                pass
        try:
            app.NormalTemplate.Saved = True
        except Exception:
            pass
        app.Quit()


def generate_delivery_document(
    template_path: Path,
    output_path: Path,
    receiver_unit: str,
    room_number: str,
    items: Iterable[ItemRow],
) -> None:
    item_rows = list(items)
    _generate_with_xml(
        template_path=template_path,
        output_path=output_path,
        unit_value=receiver_unit,
        room_number=room_number,
        items=item_rows,
        info_row_index=2,
    )


def generate_receipt_document(
    template_path: Path,
    output_path: Path,
    sender_unit: str,
    room_number: str,
    items: Iterable[ItemRow],
) -> None:
    item_rows = list(items)
    _generate_with_xml(
        template_path=template_path,
        output_path=output_path,
        unit_value=sender_unit,
        room_number=room_number,
        items=item_rows,
        info_row_index=1,
    )


def _generate_with_xml(
    *,
    template_path: Path,
    output_path: Path,
    unit_value: str,
    room_number: str,
    items: list[ItemRow],
    info_row_index: int,
) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    with zipfile.ZipFile(template_path, "r") as source_zip:
        document_xml = source_zip.read("word/document.xml")
        root = etree.fromstring(document_xml)

        tables = root.xpath(".//w:tbl", namespaces=namespaces)
        info_table = tables[0]
        items_table = tables[1]

        info_rows = info_table.xpath("./w:tr", namespaces=namespaces)
        info_cells = info_rows[info_row_index - 1].xpath("./w:tc", namespaces=namespaces)
        _set_cell_text(info_cells[1], unit_value)
        _set_cell_text(info_cells[3], room_number)

        item_rows = items_table.xpath("./w:tr", namespaces=namespaces)
        data_rows = item_rows[1:-1]
        if len(items) > len(data_rows):
            raise ValueError(
                f"O modelo aceita ate {len(data_rows)} itens, mas voce informou {len(items)}."
            )

        for offset, row in enumerate(data_rows):
            cells = row.xpath("./w:tc", namespaces=namespaces)
            patrimony = ""
            description = ""

            if offset < len(items):
                patrimony, description = items[offset]

            _set_cell_text(cells[1], patrimony)
            _set_cell_text(cells[2], description)

        updated_xml = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir=output_path.parent) as temp_file:
            temp_path = Path(temp_file.name)

        try:
            with zipfile.ZipFile(temp_path, "w") as output_zip:
                for info in source_zip.infolist():
                    data = updated_xml if info.filename == "word/document.xml" else source_zip.read(info.filename)
                    output_zip.writestr(info, data)
            temp_path.replace(output_path)
        finally:
            if temp_path.exists():
                temp_path.unlink(missing_ok=True)


def _set_cell_text(cell, text: str) -> None:
    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    tag = lambda name: f"{{{namespace}}}{name}"

    tc_pr = cell.find(tag("tcPr"))
    for child in list(cell):
        if child is not tc_pr:
            cell.remove(child)

    paragraph = etree.SubElement(cell, tag("p"))
    run = etree.SubElement(paragraph, tag("r"))
    text_element = etree.SubElement(run, tag("t"))
    if text.startswith(" ") or text.endswith(" "):
        text_element.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_element.text = text


def _fill_delivery_document(
    document: Document,
    receiver_unit: str,
    room_number: str,
    items: Iterable[ItemRow],
) -> None:
    info_table = document.tables[0]
    info_table.cell(1, 1).text = receiver_unit
    info_table.cell(1, 3).text = room_number
    _fill_items_table(document.tables[1], items)


def _fill_receipt_document(
    document: Document,
    sender_unit: str,
    room_number: str,
    items: Iterable[ItemRow],
) -> None:
    info_table = document.tables[0]
    info_table.cell(0, 1).text = sender_unit
    info_table.cell(0, 3).text = room_number
    _fill_items_table(document.tables[1], items)


def _fill_items_table(table, items: Iterable[ItemRow]) -> None:
    item_rows = list(items)
    data_start_row = 1
    available_rows = len(table.rows) - data_start_row - 1

    if len(item_rows) > available_rows:
        raise ValueError(
            f"O modelo aceita ate {available_rows} itens, mas voce informou {len(item_rows)}."
        )

    for offset in range(available_rows):
        row = table.rows[data_start_row + offset]
        patrimony = ""
        description = ""

        if offset < len(item_rows):
            patrimony, description = item_rows[offset]

        row.cells[1].text = patrimony
        row.cells[2].text = description
        row.cells[3].text = description
